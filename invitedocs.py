#! python3
# -*- coding: utf-8 -*-
#
# Solution to Practice Project "Spreadsheet Cell Inverter"
# from "Automate The Boring Stuff", by Al Sweigart,
# Chapter 12.
#
# "Say you have a text file of guest names. This guests.txt file 
# has one name per line.
# 
# Write a program that would generate a Word document with custom 
# invitations that look like Figure 13-11: 
# https://automatetheboringstuff.com/chapter13/#calibre_link-122"
# 
# Nadia Borsch      misc@nborsch.com        Jun/2018

import docx

# Set up documents and files
template = docx.Document("template.docx")
guestlist = open("guests.txt", "r")
guests = guestlist.readlines()

for guest in guests:
    # Copy text for each paragraph from template and add styles
    template.add_paragraph("It would be a pleasure to have the company of", style="Invitation")
    template.add_paragraph(guest.strip("\n"), style="Guest")
    template.add_paragraph("at 11010 Memory Lane on the evening of", style="Invitation")
    template.add_paragraph("April 1st", style="Date")
    template.add_paragraph("at 7 o’ clock", style="Invitation")

    # Page break to separate invitations
    template.add_page_break()

# Close and save
guestlist.close()
try:
    template.save("invitations.docx")
except PermissionError:
    print("Could not save invitations file, please try again.")